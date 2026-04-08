#!/usr/bin/env python3
"""Minimal docx I/O for AIGC-Detector Skill.

Commands:
  read <file>                  Print all paragraphs with numbered labels
  replace <file> <idx>          Replace paragraph <idx> (1-based) from stdin
  write <file>                 Write plain text from stdin to new .docx
"""

import sys
import os
import shutil


def read_docx(file_path: str) -> str:
    """Extract numbered plain text from a .docx file."""
    from docx import Document
    doc = Document(file_path)
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            paragraphs.append(f"[{i + 1}] {text}")
    return "\n\n".join(paragraphs)


def replace_paragraph(file_path: str, index: int, new_text: str, output_path: str = None) -> str:
    """Replace a single paragraph in-place, preserving original formatting.

    Args:
        file_path: Path to source .docx
        index: 1-based paragraph index to replace
        new_text: Replacement text
        output_path: Output path (defaults to {file}_rewritten.docx)
    """
    from docx import Document

    if output_path is None:
        base, ext = os.path.splitext(file_path)
        output_path = f"{base}_rewritten{ext}"

    # Copy to preserve original
    shutil.copy2(file_path, output_path)

    doc = Document(output_path)
    if index < 1 or index > len(doc.paragraphs):
        print(f"Error: paragraph index {index} out of range (1-{len(doc.paragraphs)})", file=sys.stderr)
        # Clean up failed copy
        if output_path != file_path:
            os.remove(output_path)
        sys.exit(1)

    para = doc.paragraphs[index - 1]
    # Preserve paragraph style
    style = para.style
    para.clear()
    run = para.add_run(new_text)
    run.font.size = para.runs[0].font.size if para.runs else None
    run.font.name = para.runs[0].font.name if para.runs else None
    para.style = style

    doc.save(output_path)
    return output_path


def write_docx(file_path: str, text: str) -> None:
    """Write plain text to a new .docx file."""
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    paragraphs = text.split("\n\n")
    for para_text in paragraphs:
        if para_text.strip():
            p = doc.add_paragraph(para_text.strip())
            for run in p.runs:
                run.font.size = Pt(12)
                run.font.name = "宋体"
    doc.save(file_path)


def main():
    if len(sys.argv) < 3:
        print("Usage: python3 docx_io.py <read|replace|write> <file_path> [args...]", file=sys.stderr)
        sys.exit(1)

    command = sys.argv[1]
    file_path = sys.argv[2]

    if command == "read":
        if not os.path.exists(file_path):
            print(f"Error: file not found: {file_path}", file=sys.stderr)
            sys.exit(1)
        text = read_docx(file_path)
        print(text)
    elif command == "replace":
        if len(sys.argv) < 4:
            print("Usage: python3 docx_io.py replace <file_path> <paragraph_index>", file=sys.stderr)
            print("  Paragraph text is read from stdin.", file=sys.stderr)
            sys.exit(1)
        index = int(sys.argv[3])
        new_text = sys.stdin.read().strip()
        output = replace_paragraph(file_path, index, new_text)
        print(output, file=sys.stderr)
    elif command == "write":
        text = sys.stdin.read()
        write_docx(file_path, text)
        print(f"Written to: {file_path}", file=sys.stderr)
    else:
        print(f"Error: unknown command '{command}'. Use 'read', 'replace', or 'write'.", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
